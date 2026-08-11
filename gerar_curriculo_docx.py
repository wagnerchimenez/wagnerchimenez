#!/usr/bin/env python3
"""Gera Curriculo_Wagner_Lima_Chimenez.docx na raiz do repositório."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Cm


OUT = Path(__file__).resolve().parent / "Curriculo_Wagner_Lima_Chimenez.docx"


def set_run_font(run, size=10, bold=False, color=None, name="Calibri"):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)
    if color is not None:
        run.font.color.rgb = color


def add_paragraph(doc, text="", *, size=10, bold=False, space_before=0, space_after=4, align=None, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_runs(p, parts):
    for text, bold, size, color in parts:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)


def add_heading_line(doc, text):
    p = add_paragraph(
        doc,
        text.upper(),
        size=11,
        bold=True,
        space_before=10,
        space_after=2,
        color=RGBColor(0x1A, 0x3A, 0x5C),
    )
    pPr = p._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "12",
            qn("w:space"): "4",
            qn("w:color"): "1A3A5C",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_job(doc, title, meta, bullets):
    p = add_paragraph(doc, "", size=10, space_before=6, space_after=0)
    add_runs(
        p,
        [
            (title, True, 10, None),
            ("  |  ", False, 10, RGBColor(0x66, 0x66, 0x66)),
            (meta, False, 9, RGBColor(0x44, 0x44, 0x44)),
        ],
    )
    for b in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(1)
        bp.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        run = bp.add_run(b)
        set_run_font(run, size=9.5)


def add_product(doc, name, url, bullets):
    p = add_paragraph(doc, "", size=10, space_before=4, space_after=0)
    add_runs(
        p,
        [
            (name, True, 10, None),
            ("  —  ", False, 10, RGBColor(0x66, 0x66, 0x66)),
            (url, False, 9, RGBColor(0x1A, 0x5C, 0x96)),
        ],
    )
    for b in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(1)
        run = bp.add_run(b)
        set_run_font(run, size=9.5)


def main():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(1.4)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.6)
        section.right_margin = Cm(1.6)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

    add_paragraph(
        doc,
        "Wagner Lima Chimenez",
        size=18,
        bold=True,
        space_after=2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x1A, 0x3A, 0x5C),
    )
    add_paragraph(
        doc,
        "Desenvolvedor Full Stack Sênior · PHP · Laravel · React",
        size=11,
        bold=True,
        space_after=2,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    add_paragraph(
        doc,
        "Dourados, MS – Brasil | Remoto | (67) 99926-2419 | wagnerllchimenez.comp@gmail.com",
        size=9,
        space_after=1,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x33, 0x33, 0x33),
    )
    add_paragraph(
        doc,
        "LinkedIn: linkedin.com/in/wagner-lima-chimenez-2b0386153  ·  GitHub: github.com/wagnerchimenez",
        size=9,
        space_after=6,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        color=RGBColor(0x33, 0x33, 0x33),
    )

    add_heading_line(doc, "Resumo")
    add_paragraph(
        doc,
        "Desenvolvedor Full Stack Sênior com ~15 anos em aplicações web, foco em PHP 8, Laravel e React (Inertia). "
        "Experiência em e-commerce de médio/grande porte (Tray), SaaS próprios em produção, APIs REST, filas (RabbitMQ), "
        "Redis, Docker, TDD e monitoramento (New Relic/Graylog). Também atuei com Symfony em sistemas corporativos. "
        "Busco oportunidades remotas como Full Stack ou Backend PHP/Laravel.",
        size=9.5,
        space_after=4,
    )

    add_heading_line(doc, "Produtos e projetos em produção")
    add_product(
        doc,
        "Officina Pro — SaaS próprio",
        "https://officinapro.com.br",
        [
            "SaaS de gestão para oficinas mecânicas em produção",
            "Stack: PHP, Laravel, Inertia.js — desenvolvimento, manutenção e evolução do produto",
        ],
    )
    add_product(
        doc,
        "Salão Elite — SaaS próprio",
        "https://salaoelite.com.br",
        [
            "SaaS de gestão para salões de beleza em produção",
            "Stack: PHP, Laravel, Inertia.js — desenvolvimento e evolução do produto",
        ],
    )
    add_product(
        doc,
        "Iguatemi Notícias — cliente",
        "https://iguateminoticias.com.br",
        [
            "Portal de notícias regionais em produção (PHP/Laravel) — desenvolvimento e manutenção",
        ],
    )

    add_heading_line(doc, "Competências técnicas")
    skills = [
        ("Backend: ", "PHP 8, Laravel, Symfony, Inertia.js, APIs REST, Composer"),
        ("Frontend: ", "React, JavaScript, TypeScript, Vue.js"),
        ("Dados / async: ", "MySQL, PostgreSQL, Redis, RabbitMQ"),
        ("Qualidade / ops: ", "PHPUnit, Pest, TDD, Docker, New Relic, Graylog, Git, Scrum, Kanban"),
        ("Complementar: ", "NestJS, React Native, Flutter"),
    ]
    for label, value in skills:
        p = add_paragraph(doc, "", size=9.5, space_before=1, space_after=1)
        add_runs(p, [(label, True, 9.5, None), (value, False, 9.5, None)])

    add_heading_line(doc, "Experiência profissional")

    add_job(
        doc,
        "Desenvolvedor Full Stack Sênior — Knet (temporário)",
        "dez/2025 – atual · Remoto",
        [
            "Evoluí sistema web com Laravel no backend e Vue.js/PrimeVue no frontend; módulos em React com TypeScript",
            "Implementei e mantive APIs REST e integração com meios de pagamento; MySQL",
            "Atuei remoto com entregas contínuas de funcionalidades e correções",
        ],
    )
    add_job(
        doc,
        "Desenvolvedor Full Stack Sênior — BDS Sistemas (temporário)",
        "ago/2025 – out/2025 · Remoto",
        [
            "Desenvolvi chat interno para suporte (NestJS + PostgreSQL); integração Flutter/React via REST e WebSocket",
            "Organizei o ambiente de desenvolvimento com containerização",
        ],
    )
    add_job(
        doc,
        "Desenvolvedor Full Stack Sênior — Tray",
        "mar/2022 – jan/2025",
        [
            "Desenvolvi funcionalidades na plataforma de e-commerce com PHP, Laravel e Vue.js",
            "Criei e integrei APIs REST para comunicação entre sistemas internos e parceiros",
            "Apliquei TDD com PHPUnit; monitei com New Relic e Graylog; filas RabbitMQ; Docker",
        ],
    )
    add_job(
        doc,
        "Desenvolvedor Web — PJBank",
        "mar/2021 – fev/2022",
        [
            "Evoluí sistema de gestão condominial com PHP, Symfony e Zend Framework",
            "Integrei APIs REST; testes com PHPUnit e Behat; New Relic; RabbitMQ; Docker",
        ],
    )
    add_job(
        doc,
        "Desenvolvedor Web — Before",
        "nov/2020 – abr/2021",
        [
            "Sistema de telefonia móvel e vendas: Laravel/Symfony, Vue.js, TypeScript, MySQL, APIs REST, PHPUnit, RabbitMQ e Docker",
        ],
    )
    add_job(
        doc,
        "Desenvolvedor Web — Missão Caiuá",
        "mar/2018 – nov/2020 · Dourados, MS",
        [
            "Sistema de RH com PHP, Laravel/Symfony, Vue.js, MySQL e APIs REST; melhorias de usabilidade no frontend",
        ],
    )
    add_job(
        doc,
        "Desenvolvedor Web — VEGA Tecnologia",
        "jul/2016 – nov/2017",
        [
            "Plataforma EAD com PHP, Laravel, Vue.js, MySQL e APIs REST",
        ],
    )
    add_job(
        doc,
        "Programador — Palancio Soluções",
        "jan/2014 – dez/2014",
        [
            "Sistema PDV (frente de caixa) com PHP/Laravel, JavaScript e MySQL",
        ],
    )
    add_job(
        doc,
        "Desenvolvedor Web — UNIGRAN",
        "jan/2011 – nov/2013",
        [
            "Sites e plataforma EAD: PHP, HTML, CSS, JavaScript e MySQL",
        ],
    )

    add_heading_line(doc, "Formação")
    add_paragraph(
        doc,
        "Pós-graduação em Gestão de TI — Anhanguera Educacional · 2013–2014",
        size=9.5,
        space_before=2,
        space_after=1,
    )
    add_paragraph(
        doc,
        "Bacharelado em Ciência da Computação — UNIGRAN · 2005–2010",
        size=9.5,
        space_before=0,
        space_after=2,
    )

    add_heading_line(doc, "Certificações relevantes")
    certs = [
        "IA for Devs Full Cycle – Full Cycle (12/2025)",
        "Vue.js 3 Composition API – Udemy (12/2025)",
        "TypeScript – Udemy (03/2025)",
        "Clean Architecture / DDD com PHP – Alura (09/2021)",
        "TDD com PHPUnit / Arquitetura Hexagonal – Udemy (06/2021)",
        "SOLID e Design Patterns em PHP – Alura (2021)",
    ]
    for c in certs:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.space_before = Pt(0)
        bp.paragraph_format.space_after = Pt(1)
        run = bp.add_run(c)
        set_run_font(run, size=9.5)

    doc.save(OUT)
    print(f"Gerado: {OUT}")


if __name__ == "__main__":
    main()
